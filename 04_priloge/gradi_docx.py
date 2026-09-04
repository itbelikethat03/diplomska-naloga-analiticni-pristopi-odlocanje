# -*- coding: utf-8 -*-
"""Sestavi Priloge.docx iz pripravljenih tabel (_podatki.json)."""
import json, os, re, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8", errors="replace")
MAPA = os.path.dirname(os.path.abspath(__file__))
D = json.load(open(os.path.join(MAPA, "_podatki.json"), encoding="utf-8"))
STEVILKA = re.compile(r"^[\d.,+\-]+( %| dni)?$")

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(11)

s = doc.sections[0]
for a in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(s, a, Cm(2.5))


def senci(celica, barva):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), barva)
    celica._tc.get_or_add_tcPr().append(el)


def naslov(besedilo):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(besedilo)
    r.bold = True
    r.font.size = Pt(12)


def opomba(besedilo):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(besedilo)
    r.italic = True
    r.font.size = Pt(9)


def vir():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    p.add_run("Vir: Lasten").font.size = Pt(9)


def slo_datum(s):
    """2024-12-23 -> 23. 12. 2024 (brez vodilnih nicel, kot v besedilu naloge)."""
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})$", str(s))
    return f"{int(m[3])}. {int(m[2])}. {m[1]}" if m else s


def ponovi_glavo(vrstica):
    tr = vrstica._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr.append(el)


def tabela(vrstice, sirine_cm, krepek_zadnji=False):
    t = doc.add_table(rows=0, cols=len(vrstice[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    zadnji = len(vrstice) - 1
    for i, v in enumerate(vrstice):
        vr = t.add_row()
        if i == 0:
            ponovi_glavo(vr)
        for j, besedilo in enumerate(v):
            besedilo = slo_datum(besedilo)
            c = vr.cells[j]
            c.width = Cm(sirine_cm[j])
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if i > 0 and j > 0 and STEVILKA.match(str(besedilo)):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(str(besedilo))
            r.font.size = Pt(9)
            if i == 0:
                r.bold = True
                senci(c, "D9D9D9")
            elif krepek_zadnji and i == zadnji:
                r.bold = True
    # sirine je treba nastaviti na vsaki celici posebej
    for vr in t.rows:
        for j, c in enumerate(vr.cells):
            c.width = Cm(sirine_cm[j])
    return t


def prelom():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# --- Priloga 1 -------------------------------------------------------------
naslov("Priloga 1: Struktura uvrstitve vrst napak v šest kategorij")
tabela(D["p1"], [5.2, 2.2, 3.0, 3.0, 2.6], krepek_zadnji=True)
opomba("Opomba: 29 oznak je bilo uvrščenih po vsebinskem pomenu pri pripravi podatkov. "
       "Pri dveh nadaljnjih oznakah iz istega obdobja (164 reklamacij) je bila kategorija "
       "zapisana že v izvornih podatkih. Izvorne oznake napak predstavljajo poslovno "
       "skrivnost podjetja in zato niso navedene.")
vir()
prelom()

# --- Priloga 2 -------------------------------------------------------------
naslov("Priloga 2: Pregled podatkovnega modela baze servis_db")
tabela(D["p2"], [2.6, 4.1, 1.7, 7.6], krepek_zadnji=True)
vir()
prelom()

# --- Priloga 3 -------------------------------------------------------------
naslov("Priloga 3: Izsek cevovoda ETL")
for vrstica in D["koda"].split("\n"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(vrstica if vrstica.strip() else " ")
    r.font.name = "Consolas"
    r.font.size = Pt(7.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
vir()
prelom()

# --- Priloga 4 -------------------------------------------------------------
naslov("Priloga 4: Izpadi objav tedenskih poročil NIJZ")
tabela(D["p4"], [1.9, 2.6, 3.3, 3.4, 4.8])
opomba("Opomba: arhiv zajema obdobje od 4. 4. 2024 do 15. 7. 2026. Merila za prepoznavo "
       "dveh nezanesljivih tednov so navedena v Tabeli 10 v poglavju 5.4.1.")
vir()

naslov("Priloga 4a: Izpeljava nabora tednov in vrst zdravstvenih storitev")
tabela(D["p4b"], [2.3, 10.7, 3.0])
vir()

# --- Priloga 5 (lezece) ----------------------------------------------------
nova = doc.add_section()
nova.orientation = WD_ORIENT.LANDSCAPE
nova.page_width, nova.page_height = nova.page_height, nova.page_width
for a in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(nova, a, Cm(2.0))

naslov("Priloga 5: Drseče preverjanje letnih napovedi po kategorijah napak")
tabela(D["p5"], [3.4, 1.4, 2.0, 2.1, 2.1, 2.2, 2.2, 3.0, 2.0])
opomba("Opomba: 48 testnih točk (6 kategorij × 8 testnih let 2017–2024). Stolpec »Učnih "
       "točk« prikazuje velikost rastočega učnega okna. Izbrani model je bil določen na "
       "zadnjem učnem obdobju, brez vpogleda v testno vrednost.")
vir()

doc.save(os.path.join(MAPA, "Priloge.docx"))
print("Priloge.docx zapisan")
